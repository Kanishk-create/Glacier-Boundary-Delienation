import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import io, os

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm, cm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame,
    Paragraph, Spacer, Table, TableStyle,
    Image, FrameBreak, PageBreak, KeepTogether, NextPageTemplate
)
from reportlab.graphics.shapes import Drawing, Rect, String, Line, Polygon
from reportlab.graphics import renderPDF

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────────────────────────────────────
W, H = A4
M = 11 * mm

# Palette
GLACIER   = colors.HexColor('#1A6FA8')
TEAL      = colors.HexColor('#0F6E56')
AMBER     = colors.HexColor('#D97706')
RED       = colors.HexColor('#B91C1C')
DARK      = colors.HexColor('#1E293B')
SNOW      = colors.HexColor('#F1F5F9')
LIGHTBLUE = colors.HexColor('#DBEAFE')
LIGHTTEAL = colors.HexColor('#D1FAE5')
LIGHTYELL = colors.HexColor('#FEF3C7')
WINNERGRN = colors.HexColor('#D1FAE5')
WHITE     = colors.white
MID       = colors.HexColor('#475569')
BORDER    = colors.HexColor('#CBD5E1')
PURPLE    = colors.HexColor('#6D28D9')
LIGHTPUR  = colors.HexColor('#EDE9FE')

IMG_DIR = '__results___files'

# ─── Styles ──────────────────────────────────────────────────────────────────
def S(name, **kw):
    base = dict(fontName='Helvetica', fontSize=7.5, leading=10.5,
                textColor=DARK, spaceAfter=2, spaceBefore=0)
    base.update(kw)
    return ParagraphStyle(name, **base)

sBody  = S('body',  leading=10,  spaceAfter=2, alignment=TA_JUSTIFY)
sBullet= S('bul',   fontSize=7,  leading=9.5,  leftIndent=8, spaceAfter=1.5)
sNote  = S('note',  fontName='Helvetica-Oblique', fontSize=6.3, textColor=MID,
           leading=8.5, spaceAfter=2)
sCapt  = S('capt',  fontName='Helvetica-Bold', fontSize=6.5, textColor=GLACIER,
           leading=8.5, spaceAfter=1, alignment=TA_CENTER)
sTH    = S('th',    fontName='Helvetica-Bold', fontSize=6.6, textColor=WHITE,
           leading=9, alignment=TA_CENTER)
sTD    = S('td',    fontSize=6.6, leading=9, alignment=TA_CENTER)
sTDL   = S('tdl',   fontSize=6.6, leading=9, alignment=TA_LEFT)
sTDBold= S('tdb',   fontName='Helvetica-Bold', fontSize=6.6, textColor=TEAL,
           leading=9, alignment=TA_CENTER)
sSecH  = S('sech',  fontName='Helvetica-Bold', fontSize=7.2, textColor=WHITE,
           leading=10, spaceAfter=0)

# ─── Helpers ─────────────────────────────────────────────────────────────────
def sec(title, color=GLACIER, w=None, icon='▪'):
    cw = w or (W - 2*M)
    t = Table([[Paragraph(f'{icon}  {title}', sSecH)]],
              colWidths=[cw])
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), color),
        ('TOPPADDING',    (0,0), (-1,-1), 3.5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3.5),
        ('LEFTPADDING',   (0,0), (-1,-1), 7),
        ('RIGHTPADDING',  (0,0), (-1,-1), 4),
    ]))
    return t

def SP(n=2): return Spacer(1, n)

def img(fname, w, h=None, caption=None):
    path = os.path.join(IMG_DIR, fname)
    ratio = _img_ratio(path)
    ih = h if h else w / ratio
    items = [Image(path, width=w, height=ih)]
    if caption:
        items.append(SP(2))
        items.append(Paragraph(caption, sCapt))
    return items

def _img_ratio(path):
    from PIL import Image as PILImage
    im = PILImage.open(path)
    return im.size[0] / im.size[1]

def metric_strip(items, total_w):
    n = len(items)
    cw = total_w / n
    row1 = [Paragraph(
        f'<font name="Helvetica-Bold" size="13" color="#1A6FA8">{v}</font>',
        S('mv', alignment=TA_CENTER, leading=15, spaceAfter=0))
        for v, _ in items]
    row2 = [Paragraph(
        f'<font name="Helvetica" size="5.8" color="#64748B">{l}</font>',
        S('ml', alignment=TA_CENTER, leading=8, spaceAfter=0))
        for _, l in items]
    t = Table([row1, row2], colWidths=[cw]*n)
    t.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), SNOW),
        ('TOPPADDING',    (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING',   (0,0), (-1,-1), 2),
        ('RIGHTPADDING',  (0,0), (-1,-1), 2),
        ('LINEAFTER',     (0,0), (-2,-1), 0.4, BORDER),
    ]))
    return t

def grid_table(cell_list, ncols, col_w, spacing=3):
    """Wrap a flat list of flowables into a grid table"""
    rows = []
    for i in range(0, len(cell_list), ncols):
        rows.append(cell_list[i:i+ncols])
    # pad last row
    while len(rows[-1]) < ncols:
        rows[-1].append('')
    t = Table(rows, colWidths=[col_w]*ncols,
              hAlign='LEFT',
              style=TableStyle([
                  ('TOPPADDING',    (0,0), (-1,-1), spacing),
                  ('BOTTOMPADDING', (0,0), (-1,-1), spacing),
                  ('LEFTPADDING',   (0,0), (-1,-1), spacing),
                  ('RIGHTPADDING',  (0,0), (-1,-1), spacing),
                  ('VALIGN',        (0,0), (-1,-1), 'TOP'),
              ]))
    return t

# ─── Architecture drawing ─────────────────────────────────────────────────────
def arch_drawing(w, h=105):
    d = Drawing(w, h)
    bh = 12; pad = 2
    gap = 6
    n_rows = 5
    total_h = n_rows * bh + (n_rows-1) * gap + 14  # 14 for group label
    # rows from bottom
    r5 = 4;           r4 = r5+bh+gap; r3 = r4+bh+gap
    r2 = r3+bh+gap;   r1 = r2+bh+gap

    def box(x, y, bw, bht, lbl, sub='', fc=LIGHTBLUE, tc=GLACIER, fs=5.8):
        d.add(Rect(x, y, bw, bht, rx=3, ry=3,
                   fillColor=fc, strokeColor=BORDER, strokeWidth=0.4))
        mx = x + bw/2
        if sub:
            d.add(String(mx, y+bht/2+1.5, lbl,
                         fontName='Helvetica-Bold', fontSize=fs,
                         fillColor=tc, textAnchor='middle'))
            d.add(String(mx, y+bht/2-4.5, sub,
                         fontName='Helvetica', fontSize=4.8,
                         fillColor=MID, textAnchor='middle'))
        else:
            d.add(String(mx, y+bht/2-2.5, lbl,
                         fontName='Helvetica-Bold', fontSize=fs,
                         fillColor=tc, textAnchor='middle'))

    def arr(x1,y1,x2,y2):
        d.add(Line(x1,y1,x2,y2, strokeColor=MID, strokeWidth=0.6))
        dx,dy=x2-x1,y2-y1; ln=(dx**2+dy**2)**0.5
        if ln<1: return
        ux,uy=dx/ln,dy/ln
        d.add(Polygon([x2,y2, x2-3.5*ux+2*uy, y2-3.5*uy-2*ux,
                              x2-3.5*ux-2*uy, y2-3.5*uy+2*ux],
                      fillColor=MID, strokeColor=MID, strokeWidth=0.3))

    def grp(x,y,bw,bht,lbl,fc):
        d.add(Rect(x,y,bw,bht, rx=4, ry=4,
                   fillColor=colors.Color(fc.red,fc.green,fc.blue,0.07),
                   strokeColor=fc, strokeWidth=0.7, strokeDashArray=[3,2]))
        d.add(String(x+4, y+bht-7, lbl,
                     fontName='Helvetica-Bold', fontSize=5,
                     fillColor=fc, textAnchor='start'))

    bw3 = (w-4*pad)/3
    hw  = (w-3*pad)/2

    # Row 1
    box(pad,           r1, bw3, bh, 'Sentinel-1 SAR','Radar',  LIGHTBLUE, GLACIER)
    box(pad+bw3+pad,   r1, bw3, bh, 'Sentinel-2',    'Optical',LIGHTBLUE, GLACIER)
    box(pad+2*(bw3+pad),r1, bw3, bh,'EnMAP 218-band','Hyper',  LIGHTBLUE, GLACIER)
    # Row 2
    box(pad,           r2, bw3, bh, 'downloader.py', '', LIGHTTEAL, TEAL)
    box(pad+bw3+pad,   r2, bw3, bh, 'processor.py',  '', LIGHTTEAL, TEAL)
    box(pad+2*(bw3+pad),r2, bw3, bh,'MongoDB',       '', LIGHTTEAL, TEAL)
    # Row 3
    box(pad,           r3, hw, bh, 'U-Net Seg.','IoU 0.938', LIGHTYELL, AMBER)
    box(pad+hw+pad,    r3, hw-pad, bh,'ML Classify','99.68% BA', LIGHTYELL, AMBER)
    # Row 4
    box(pad,           r4, hw, bh, 'FastAPI Backend','REST/WS', LIGHTBLUE, GLACIER)
    box(pad+hw+pad,    r4, hw-pad, bh,'GLOF Alerts','change_detector',
        colors.HexColor('#FEE2E2'), RED)
    # Row 5
    bw5=(w-4*pad)/3
    box(pad,           r5, bw5, bh, 'Mobile App','React Native',LIGHTBLUE,GLACIER)
    box(pad+bw5+pad,   r5, bw5, bh, 'REST API',  'Data license',LIGHTBLUE,GLACIER)
    box(pad+2*(bw5+pad),r5,bw5, bh, 'Dashboard', 'Web UI',      LIGHTBLUE,GLACIER)

    # Groups
    grp(0,r1-2,w,bh+4,'Data Sources', GLACIER)
    grp(0,r2-2,w,bh+4,'Ingestion',    TEAL)
    grp(0,r3-2,w,bh+4,'AI Core',      AMBER)
    grp(0,r4-2,w,bh+4,'Backend',      RED)
    grp(0,r5-2,w,bh+4,'Outputs',      GLACIER)

    # Arrows
    cx1=pad+bw3/2; cx2=pad+bw3+pad+bw3/2; cx3=pad+2*(bw3+pad)+bw3/2
    arr(cx1,r1,cx1,r2+bh); arr(cx2,r1,cx2,r2+bh); arr(cx3,r1,cx3,r2+bh)
    chw1=pad+hw/2; chw2=pad+hw+pad+(hw-pad)/2
    arr(cx1,r2,chw1,r3+bh); arr(cx2,r2,chw1,r3+bh); arr(cx3,r2,chw2,r3+bh)
    arr(chw1,r3,chw1,r4+bh); arr(chw2,r3,chw2,r4+bh)
    cbs1=pad+bw5/2; cbs2=pad+bw5+pad+bw5/2; cbs3=pad+2*(bw5+pad)+bw5/2
    arr(chw1,r4,cbs1,r5+bh); arr(chw1,r4,cbs2,r5+bh); arr(chw1,r4,cbs3,r5+bh)
    return d

# ─── Model bar chart ──────────────────────────────────────────────────────────
def make_model_chart(w_pts, h_pts):
    models  = ['LDA','QDA','GBDT\nOrig','k-NN\nOrig','RF\nOrig',
               'SVM\nPCA','GBDT\nPCA','RF\nPCA','Logistic\nOrig']
    orig    = [91.65,92.07,97.07,98.68,99.14,99.31,98.97,98.59,99.68]
    pca     = [92.27,95.00,98.85,98.76,98.73,99.56,98.97,98.73,99.67]
    x       = np.arange(len(models)); width = 0.37
    dpi     = 150
    fw, fh  = w_pts/dpi*1.05, h_pts/dpi*1.05

    fig, ax = plt.subplots(figsize=(fw, fh))
    fig.patch.set_facecolor('#F8FAFC'); ax.set_facecolor('#F8FAFC')
    b1 = ax.bar(x-width/2, orig, width, label='Original features',
                color='#1A6FA8', alpha=0.88, zorder=3)
    b2 = ax.bar(x+width/2, pca,  width, label='PCA features',
                color='#0F6E56', alpha=0.75, zorder=3)
    b1[-1].set_color('#D97706'); b1[-1].set_alpha(1.0)
    ax.set_xticks(x); ax.set_xticklabels(models, fontsize=6.2)
    ax.set_ylim(88,101.5)
    ax.set_ylabel('Avg. Balanced Accuracy (%)', fontsize=6.5)
    ax.set_title('All Models — Avg. Balanced Accuracy (8 Classes)',
                 fontsize=7.2, fontweight='bold', pad=4)
    ax.tick_params(axis='y', labelsize=6.5)
    ax.yaxis.grid(True, alpha=0.35, zorder=0); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.legend(fontsize=6, loc='lower right', framealpha=0.85,
              edgecolor='#CBD5E1')
    ax.annotate('★ Best', xy=(8-width/2,99.68), xytext=(6.8,100.5),
                fontsize=6.2, color='#D97706', fontweight='bold',
                arrowprops=dict(arrowstyle='->', color='#D97706', lw=0.8))
    plt.tight_layout(pad=0.4)
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                facecolor='#F8FAFC')
    plt.close(); buf.seek(0)
    return buf

# ─── Snow/ice F1 chart ────────────────────────────────────────────────────────
def make_snowice_chart(w_pts, h_pts):
    models = ['LDA\nOrig','LDA\nPCA','QDA\nOrig','QDA\nPCA',
              'k-NN\nOrig','k-NN\nPCA','GBDT\nOrig','GBDT\nPCA',
              'RF\nOrig','RF\nPCA','SVM\nOrig','SVM\nPCA',
              'Log\nOrig','Log\nPCA']
    f1 = [0.9312,0.9112,0.9280,0.9432,
          0.9916,0.9915,0.9854,0.9918,
          0.9889,0.9897,0.9860,0.9947,
          0.9949,0.9944]
    clrs = ['#94A3B8']*12+['#D97706','#1A6FA8']
    dpi  = 150
    fw, fh = w_pts/dpi*1.05, h_pts/dpi*1.05
    fig, ax = plt.subplots(figsize=(fw, fh))
    fig.patch.set_facecolor('#F8FAFC'); ax.set_facecolor('#F8FAFC')
    ax.bar(range(len(models)), f1, color=clrs, alpha=0.88, zorder=3)
    ax.set_xticks(range(len(models))); ax.set_xticklabels(models, fontsize=5.6)
    ax.set_ylim(0.88,1.01)
    ax.set_ylabel('F1 Score', fontsize=6.5)
    ax.set_title('Snow / Ice Class (Class 4) — F1 Score by Model',
                 fontsize=7.2, fontweight='bold', pad=4)
    ax.tick_params(axis='y', labelsize=6.5)
    ax.yaxis.grid(True, alpha=0.35, zorder=0); ax.set_axisbelow(True)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.axhline(y=0.99, color='#D97706', linestyle='--', lw=0.9, alpha=0.7)
    from matplotlib.patches import Patch
    ax.legend(handles=[
        Patch(color='#D97706', label='Best (Logistic Orig)'),
        Patch(color='#1A6FA8', label='Logistic PCA'),
        Patch(color='#94A3B8', label='Other models'),
    ], fontsize=5.8, loc='lower right', framealpha=0.85, edgecolor='#CBD5E1')
    plt.tight_layout(pad=0.4)
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                facecolor='#F8FAFC')
    plt.close(); buf.seek(0)
    return buf

# ─── Page callbacks ──────────────────────────────────────────────────────────
def page1_bg(canvas, doc):
    canvas.saveState()
    hh = 34*mm
    # header
    canvas.setFillColor(GLACIER)
    canvas.rect(0, H-hh, W, hh, fill=1, stroke=0)
    # subtle dot pattern
    canvas.setFillColor(colors.Color(1,1,1,0.04))
    for xi in range(0,int(W),16):
        for yi in range(int(H-hh),int(H),13):
            canvas.circle(xi+8, yi+7, 3.5, fill=1, stroke=0)
    # Title
    canvas.setFont('Helvetica-Bold', 20)
    canvas.setFillColor(WHITE)
    canvas.drawString(M, H-15*mm, '❄  GlacioWatch')
    # page badge
    canvas.setFont('Helvetica-Bold', 7)
    canvas.setFillColor(colors.Color(1,1,1,0.25))
    canvas.roundRect(W-M-24, H-14*mm, 22, 9, 3, fill=1, stroke=0)
    canvas.setFillColor(WHITE)
    canvas.drawCentredString(W-M-13, H-12.5*mm, 'PAGE 1')
    # subtitle
    canvas.setFont('Helvetica', 8.5)
    canvas.setFillColor(LIGHTBLUE)
    canvas.drawString(M, H-21.5*mm,
        'AI-Powered Himalayan Glacier Monitoring & Early Warning System')
    # tags
    canvas.setFont('Helvetica', 6.8)
    tags = ['U-Net Segmentation','Hyperspectral ML','Sentinel-1/2 SAR',
            'GLOF Early Warning','Mobile Dashboard','218-Band EnMAP']
    tx = M
    for tag in tags:
        tw = canvas.stringWidth(tag,'Helvetica',6.8)
        canvas.setFillColor(colors.Color(1,1,1,0.17))
        canvas.roundRect(tx-3, H-28.5*mm, tw+6, 8, 3, fill=1, stroke=0)
        canvas.setFillColor(WHITE)
        canvas.drawString(tx, H-27.8*mm, tag)
        tx += tw + 10
    # author
    canvas.setFont('Helvetica', 6.8)
    canvas.setFillColor(LIGHTBLUE)
    for i,(l) in enumerate(['PVGCOE-SSDIOM, Nashik',
                             'Kapadnis J. | Pandey J. | Vadge K. | Thok P.',
                             'yash.vadge04@gmail.com']):
        canvas.drawRightString(W-M, H-15*mm+(-7*i), l)
    # column divider
    mid = W/2
    canvas.setStrokeColor(BORDER)
    canvas.setLineWidth(0.4)
    canvas.setDash([3,3])
    canvas.line(mid, 11*mm, mid, H-hh-2*mm)
    canvas.setDash()
    # footer
    canvas.setFillColor(DARK)
    canvas.rect(0, 0, W, 10*mm, fill=1, stroke=0)
    canvas.setFont('Helvetica', 6)
    canvas.setFillColor(WHITE)
    canvas.drawString(M, 3.8*mm,
        'GlacioWatch  •  U-Net IoU 0.938  •  Logistic Regression 99.68% BA  •  PVGCOE-SSDIOM, Nashik')
    canvas.drawRightString(W-M, 3.8*mm, '❄  Page 1 of 2')
    canvas.restoreState()

def page2_bg(canvas, doc):
    canvas.saveState()
    hh = 18*mm
    canvas.setFillColor(DARK)
    canvas.rect(0, H-hh, W, hh, fill=1, stroke=0)
    canvas.setFont('Helvetica-Bold', 13)
    canvas.setFillColor(WHITE)
    canvas.drawString(M, H-11*mm, '❄  GlacioWatch  —  Visual Analysis & Results (Page 2 of 2)')
    canvas.setFont('Helvetica', 7)
    canvas.setFillColor(colors.HexColor('#94A3B8'))
    canvas.drawRightString(W-M, H-11*mm,
        'Hyperspectral ML Classification  •  Tyrol Alpine Region, Austria')
    # subtle stripe
    canvas.setFillColor(GLACIER)
    canvas.rect(0, H-hh, 5, hh, fill=1, stroke=0)
    # footer
    canvas.setFillColor(DARK)
    canvas.rect(0, 0, W, 10*mm, fill=1, stroke=0)
    canvas.setFont('Helvetica', 6)
    canvas.setFillColor(WHITE)
    canvas.drawString(M, 3.8*mm,
        'GlacioWatch  •  EnMAP 218-band Hyperspectral  •  215,604 pixels  •  8 terrain classes  •  PVGCOE-SSDIOM, Nashik')
    canvas.drawRightString(W-M, 3.8*mm, '❄  Page 2 of 2')
    canvas.restoreState()

def build_docx(out):
    doc = Document()
    
    # ── Title ──
    title = doc.add_heading('GlacioWatch Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('AI-Powered Himalayan Glacier Monitoring & Early Warning System').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ── Summary ──
    doc.add_heading('Project Summary', level=1)
    doc.add_paragraph(
        "India's 9,500 Himalayan glaciers feed the Indus, Ganga & Brahmaputra. "
        "GlacioWatch combines U-Net deep segmentation on Sentinel SAR imagery "
        "with hyperspectral ML classification on 218-band EnMAP data."
    )
    
    # ── Metrics ──
    doc.add_heading('Key Performance Metrics', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'U-Net Acc'
    hdr_cells[1].text = 'U-Net IoU'
    hdr_cells[2].text = 'Logistic BA'
    hdr_cells[3].text = 'Snow/Ice F1'
    
    row_cells = table.add_row().cells
    row_cells[0].text = '97.74%'
    row_cells[1].text = '0.938'
    row_cells[2].text = '99.68%'
    row_cells[3].text = '0.9949'
    
    # ── Visual Results ──
    doc.add_heading('Visual Analysis & Results', level=1)
    
    imgs = [
        ('__results___23_0.png', 'Fig 1: Spatial Distribution of Land Types & Classes'),
        ('__results___28_0.png', 'Fig 2: Hyperspectral RGB vs Land Type Classification Map'),
        ('__results___31_0.png', 'Fig 3: 3D Spatial-Spectral Reflectance Visualization'),
        ('__results___30_0.png', 'Fig 4: Detailed Hyperspectral Land Cover Map (Tyrol)'),
        ('result_01.png',        'Fig 5: Hardangerjøkulen - Optical, Elevation & Slope Inputs'),
        ('result_02.png',        'Fig 6: Glacier Outlines & Segmentation Confidence Map'),
        ('__results___37_0.png', 'Fig 7: Spectral Reflectance Box Plots across Classes'),
        ('__results___21_0.png', 'Fig 8: Statistical Class Distribution Analysis')
    ]
    
    for fname, caption in imgs:
        path = os.path.join(IMG_DIR, fname)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(5))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph() # Spacer
            
    doc.save(out)

# ─── BUILD ───────────────────────────────────────────────────────────────────
def build():
    out = 'GlacioWatch_Report_2Page.pdf'
    out_docx = 'GlacioWatch_Report.docx'
    
    out_dir = os.path.dirname(out)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    # 1. PDF Generation
    doc = BaseDocTemplate(out, pagesize=A4,
                          leftMargin=M, rightMargin=M,
                          topMargin=36*mm, bottomMargin=13*mm)
    
    body_w = W - 2*M
    col_w  = (body_w - 6*mm) / 2
    gutter = 6*mm

    # PAGE 1 — two columns
    fL1 = Frame(M,              13*mm, col_w, H-36*mm-13*mm, id='p1L',
                leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    fR1 = Frame(M+col_w+gutter, 13*mm, col_w, H-36*mm-13*mm, id='p1R',
                leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)

    # PAGE 2 — full width
    fFull = Frame(M, 13*mm, body_w, H-18*mm-13*mm, id='p2F',
                  leftPadding=0, rightPadding=0, topPadding=3*mm, bottomPadding=0)

    doc.addPageTemplates([
        PageTemplate(id='page1', frames=[fL1, fR1], onPage=page1_bg),
        PageTemplate(id='page2', frames=[fFull],    onPage=page2_bg),
    ])

    story = []

    # ═══════════════════════════════════════════════════════════
    # PAGE 1 — LEFT COLUMN
    # ═══════════════════════════════════════════════════════════

    # ── Key metrics strip ────────────────────────────────────
    story.append(metric_strip([
        ('97.74%','U-Net Accuracy'),
        ('0.938',  'U-Net IoU'),
        ('99.68%', 'Logistic BA'),
        ('0.9949', 'Snow/Ice F1'),
    ], col_w))
    story.append(SP(5))

    # ── Crisis ───────────────────────────────────────────────
    story.append(sec('THE CRISIS', RED, col_w, '▲'))
    story.append(SP(3))
    story.append(Paragraph(
        "India's <b>9,500 Himalayan glaciers</b> (33,000 km²) feed the Indus, Ganga & "
        "Brahmaputra — lifelines for <b>700M+ people</b> and a multi-trillion-rupee "
        "agricultural and energy economy. They are disappearing at an alarming rate.", sBody))
    story.append(SP(4))

    ct = Table([
        [Paragraph('<b>Indicator</b>',sTH), Paragraph('<b>Impact</b>',sTH)],
        [Paragraph('30%+ glacier loss (40 yrs)',sTDL), Paragraph('Water & food insecurity',sTDL)],
        [Paragraph('Gangotri retreats 20 m/yr', sTDL), Paragraph('Ganga source threatened',sTDL)],
        [Paragraph('16.3% hydro output drop',   sTDL), Paragraph('India FY 2023–24',sTDL)],
        [Paragraph('195 high-risk glacial lakes',sTDL), Paragraph('NDMA 2024 assessment',sTDL)],
        [Paragraph('Sikkim GLOF Oct 2023',      sTDL), Paragraph('Teesta-III dam destroyed',sTDL)],
    ], colWidths=[col_w*0.52, col_w*0.48])
    ct.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),GLACIER),
        ('FONTSIZE',  (0,0),(-1,-1),6.6),
        ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),2),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),2),
        ('GRID',(0,0),(-1,-1),0.3,BORDER),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,SNOW]),
    ]))
    story.append(ct)
    story.append(SP(3))
    story.append(Paragraph(
        '<font color="#B91C1C"><b>Critical gap:</b></font> The 2023 Sikkim GLOF first '
        'alert came from a soldier — not a sensor. Every $1 in early warning systems '
        'averts <b>$4–$10</b> in disaster losses.', sBody))
    story.append(SP(5))

    # ── Solution ─────────────────────────────────────────────
    story.append(sec('SOLUTION: TWO-LAYER AI ENGINE', TEAL, col_w))
    story.append(SP(3))
    story.append(Paragraph(
        'GlacioWatch combines <b>U-Net deep segmentation</b> on Sentinel SAR imagery '
        'with <b>hyperspectral ML classification</b> on 218-band EnMAP data, delivering '
        'complete real-time glacier intelligence.', sBody))
    story.append(SP(3))
    for b in [
        '<b>Debris-covered ice & shadows</b> handled — where NDSI fails',
        '<b>218-band EnMAP</b> hyperspectral processing with PCA reduction',
        '<b>8 terrain classes</b>: meadow, tundra, bare/dark rock, snow/ice, scree, valley, veg-scree',
        '<b>15-min Sentinel-1 SAR</b> polling — cloud-penetrating radar',
        '<b>Mobile-first GLOF alerts</b> via React Native for field officers',
    ]:
        story.append(Paragraph(f'<font color="#1A6FA8">▸</font>  {b}', sBullet))
    story.append(SP(5))

    # ── Segmentation results ──────────────────────────────────
    story.append(sec('U-NET SEGMENTATION RESULTS', GLACIER, col_w))
    story.append(SP(3))
    st = Table([
        [Paragraph('<b>Metric</b>',sTH),    Paragraph('<b>U-Net</b>',sTH),  Paragraph('<b>CNN</b>',sTH)],
        [Paragraph('Accuracy',sTDL),        Paragraph('<b>97.74%</b>',sTDBold), Paragraph('93.24%',sTD)],
        [Paragraph('Dice Coeff.',sTDL),     Paragraph('<b>0.968</b>',sTDBold),  Paragraph('0.912',sTD)],
        [Paragraph('IoU',sTDL),             Paragraph('<b>0.938</b>',sTDBold),  Paragraph('0.838',sTD)],
        [Paragraph('Precision (Gl.)',sTDL), Paragraph('<b>0.976</b>',sTDBold),  Paragraph('0.895',sTD)],
        [Paragraph('Recall (Gl.)',sTDL),    Paragraph('<b>0.961</b>',sTDBold),  Paragraph('0.931',sTD)],
    ], colWidths=[col_w*0.46, col_w*0.27, col_w*0.27])
    st.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),GLACIER),
        ('BACKGROUND',(1,1),(1,-1),WINNERGRN),
        ('FONTSIZE',  (0,0),(-1,-1),6.6),
        ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),2),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),2),
        ('GRID',(0,0),(-1,-1),0.3,BORDER),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,SNOW]),
        ('ALIGN',(1,0),(-1,-1),'CENTER'),
    ]))
    story.append(st)
    story.append(SP(2))
    story.append(Paragraph(
        'Encoder-decoder U-Net with multi-head self-attention bottleneck. '
        'Trained on Glacier-Lake PNG dataset (70/20/10 split), 50 epochs, Adam optimizer.', sNote))
    story.append(SP(5))

    # ── Revenue ───────────────────────────────────────────────
    story.append(sec('REVENUE MODEL', AMBER, col_w, 'INR'))
    story.append(SP(3))
    rt = Table([
        [Paragraph('<b>Stream</b>',sTH), Paragraph('<b>Customers</b>',sTH), Paragraph('<b>ARR</b>',sTH)],
        [Paragraph('Govt. Contracts',sTDL), Paragraph('NDMA, CWC, SDMAs',sTDL), Paragraph('INR 50L–5Cr',sTD)],
        [Paragraph('Hydro SaaS',sTDL),      Paragraph('NHPC, SJVN, THDC',sTDL), Paragraph('INR 10–40L',sTD)],
        [Paragraph('API Licensing',sTDL),   Paragraph('Reinsurers, ESG',sTDL),   Paragraph('INR 2–15L',sTD)],
        [Paragraph('Research/NGO',sTDL),    Paragraph('IITs, ICIMOD, WWF',sTDL), Paragraph('INR 1–5L',sTD)],
    ], colWidths=[col_w*0.3, col_w*0.42, col_w*0.28])
    rt.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),AMBER),
        ('FONTSIZE',  (0,0),(-1,-1),6.6),
        ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),2),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),2),
        ('GRID',(0,0),(-1,-1),0.3,BORDER),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[WHITE,SNOW]),
    ]))
    story.append(rt)
    story.append(SP(2))
    story.append(Paragraph(
        '<b>Year 1 target:</b> 1 govt pilot + 2 hydro subscriptions = '
        '<font color="#0F6E56"><b>INR 70–90L ARR</b></font> | '
        '<b>$20M</b> NDMA GLOF programme actively seeking tech vendors.', sBody))
    
    img_path_side = os.path.join(IMG_DIR, '__results___23_0.png')
    if os.path.exists(img_path_side):
        img2_h = col_w / _img_ratio(img_path_side)
        story.append(Image(img_path_side, width=col_w, height=img2_h))
        cap2 = Paragraph('Fig 1: Spatial Distribution of Land Types', S('cap',fontSize=5.8,alignment=TA_CENTER))
        story.append(cap2)
        story.append(SP(4))

    # ── COLUMN BREAK ─────────────────────────────────────────
    story.append(FrameBreak())

    # ═══════════════════════════════════════════════════════════
    # PAGE 1 — RIGHT COLUMN
    # ═══════════════════════════════════════════════════════════

    # ── Architecture ─────────────────────────────────────────
    story.append(sec('SYSTEM ARCHITECTURE', GLACIER, col_w))
    story.append(SP(2))
    story.append(arch_drawing(col_w, h=105))
    story.append(SP(2))
    story.append(Paragraph(
        'Two AI branches (segmentation + hyperspectral classification) feed '
        'a unified FastAPI backend delivering real-time GLOF alerts via REST API '
        'and React Native mobile app.', sNote))
    story.append(SP(5))

    # ── ML comparison chart ───────────────────────────────────
    story.append(sec('ML MODEL COMPARISON — ALL 7 MODELS', TEAL, col_w))
    story.append(SP(2))
    c1 = make_model_chart(col_w * 2.835, col_w * 2.835 * 0.35)
    story.append(Image(c1, width=col_w, height=col_w*0.35))
    story.append(SP(2))
    story.append(Paragraph(
        'Logistic Regression (Original 218-band) achieves highest balanced '
        'accuracy (99.68%). PCA variants competitive. LDA/QDA lag significantly.', sNote))
    story.append(SP(5))

    # ── Per-class table ───────────────────────────────────────
    story.append(sec('ALL MODELS — AVG. PER-CLASS METRICS', GLACIER, col_w))
    story.append(SP(3))
    ml = Table([
        [Paragraph('<b>Model</b>',sTH), Paragraph('<b>Bal. Acc.</b>',sTH),
         Paragraph('<b>F1</b>',sTH),   Paragraph('<b>AUC</b>',sTH)],
        [Paragraph('<b>Logistic (Orig)</b>',sTDBold), Paragraph('<b>99.68%</b>',sTDBold),
         Paragraph('<b>0.9941</b>',sTDBold),          Paragraph('<b>0.9999</b>',sTDBold)],
        [Paragraph('Logistic (PCA)',sTDL),  Paragraph('99.67%',sTD), Paragraph('0.9928',sTD), Paragraph('0.9999',sTD)],
        [Paragraph('SVM (PCA)',sTDL),       Paragraph('99.56%',sTD), Paragraph('0.9904',sTD), Paragraph('—',sTD)],
        [Paragraph('RF (Orig)',sTDL),       Paragraph('99.14%',sTD), Paragraph('0.9822',sTD), Paragraph('0.9999',sTD)],
        [Paragraph('GBDT (PCA)',sTDL),      Paragraph('98.97%',sTD), Paragraph('0.9812',sTD), Paragraph('0.9997',sTD)],
        [Paragraph('k-NN (Orig)',sTDL),     Paragraph('98.68%',sTD), Paragraph('0.9773',sTD), Paragraph('0.9990',sTD)],
        [Paragraph('QDA (PCA)',sTDL),       Paragraph('94.99%',sTD), Paragraph('0.9110',sTD), Paragraph('0.9964',sTD)],
        [Paragraph('LDA (Orig)',sTDL),      Paragraph('91.65%',sTD), Paragraph('0.8602',sTD), Paragraph('0.9906',sTD)],
    ], colWidths=[col_w*0.37, col_w*0.21, col_w*0.21, col_w*0.21])
    ml.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),GLACIER),
        ('BACKGROUND',(0,1),(-1,1),WINNERGRN),
        ('FONTSIZE',  (0,0),(-1,-1),6.5),
        ('TOPPADDING',(0,0),(-1,-1),2),('BOTTOMPADDING',(0,0),(-1,-1),2),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),2),
        ('GRID',(0,0),(-1,-1),0.3,BORDER),
        ('ROWBACKGROUNDS',(0,2),(-1,-1),[WHITE,SNOW]),
        ('ALIGN',(1,0),(-1,-1),'CENTER'),
    ]))
    story.append(ml)
    story.append(SP(2))
    story.append(Paragraph(
        '<font color="#B91C1C"><b>Snow/Ice (Class 4) — Logistic Orig:</b></font> '
        'BA=99.78%, F1=0.9949, AUC=0.9999 — near-perfect glacier-class isolation.', sNote))
    story.append(SP(5))

    # ── Snow/ice F1 chart ─────────────────────────────────────
    story.append(sec('SNOW / ICE CLASS — F1 BY MODEL', RED, col_w, '❄'))
    story.append(SP(2))
    c2 = make_snowice_chart(col_w*2.835, col_w*2.835*0.32)
    story.append(Image(c2, width=col_w, height=col_w*0.32))
    story.append(SP(2))
    story.append(Paragraph(
        'Logistic (Orig) & SVM (PCA) top glacier-class F1. LDA/QDA (Orig) '
        'confirm linear boundaries are insufficient without PCA for '
        'hyperspectral terrain discrimination.', sNote))
    story.append(SP(5))

    # ── Tech stack ────────────────────────────────────────────
    story.append(sec('TECHNOLOGY STACK', DARK, col_w))
    story.append(SP(3))
    tech = [
        ('Segmentation',   'U-Net + MHSA, Custom CNN'),
        ('Classification', 'Logistic, RF, SVM, k-NN, GBDT, LDA, QDA'),
        ('Features',       '218-band EnMAP + PCA (10 components)'),
        ('Imagery',        'Sentinel-1 SAR, Sentinel-2, Landsat'),
        ('Backend',        'FastAPI (Python), MongoDB'),
        ('Frontend',       'React Native + Expo')
    ]
    tt = Table([[Paragraph(f'<b>{k}</b>',S('tk',fontSize=6.3,textColor=GLACIER)),
                 Paragraph(v, S('tv',fontSize=6.3))] for k,v in tech],
               colWidths=[col_w*0.34, col_w*0.66])
    tt.setStyle(TableStyle([
        ('TOPPADDING',(0,0),(-1,-1),2.5),('BOTTOMPADDING',(0,0),(-1,-1),2.5),
        ('LEFTPADDING',(0,0),(-1,-1),4),('RIGHTPADDING',(0,0),(-1,-1),3),
        ('LINEBELOW',(0,0),(-1,-2),0.3,BORDER),
        ('ROWBACKGROUNDS',(0,0),(-1,-1),[WHITE,SNOW]),
    ]))
    story.append(tt)

    # ── Page 2 ────────────────────────────────────────────────

    story.append(PageBreak())

    story.append(sec('DATASET & STUDY REGION', TEAL, col_w))
    story.append(SP(4))
    story.append(metric_strip([
        ('215,604', 'Total Pixels'),    
        ('218',     'Spectral Bands'),
        ('8',       'Terrain Classes'),
        ('420–2450nm','Wavelength Range'),
    ], col_w))
    story.append(SP(6))

    # --- Figure 2: RGB vs Classification ---
    img_path_2 = os.path.join(IMG_DIR, '__results___28_0.png')
    if os.path.exists(img_path_2):
        img_h = col_w / _img_ratio(img_path_2)
        story.append(Image(img_path_2, width=col_w, height=img_h))
        story.append(Paragraph('Fig 2: Hyperspectral RGB vs Land Type Classification Map', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))
    
    # --- Figure 3: 3D Visualization ---
    img_path_3 = os.path.join(IMG_DIR, '__results___31_0.png')
    if os.path.exists(img_path_3):
        img3_h = col_w / _img_ratio(img_path_3)
        story.append(Image(img_path_3, width=col_w, height=img3_h))
        story.append(Paragraph('Fig 3: 3D Spatial-Spectral Reflectance Visualization', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))
    
    # --- Figure 4: Hardangerjøkulen Inputs ---
    img_path_4 = os.path.join(IMG_DIR, 'result_01.png')
    if os.path.exists(img_path_4):
        img4_h = col_w / _img_ratio(img_path_4)
        story.append(Image(img_path_4, width=col_w, height=img4_h))
        story.append(Paragraph('Fig 4: Hardangerjøkulen - Optical, Elevation & Slope Inputs', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))

    # --- Figure 5: Spectral Box Plots ---
    img_path_5 = os.path.join(IMG_DIR, '__results___37_0.png')
    if os.path.exists(img_path_5):
        img5_h = col_w / _img_ratio(img_path_5)
        story.append(Image(img_path_5, width=col_w, height=img5_h))
        story.append(Paragraph('Fig 5: Spectral Reflectance Box Plots across classes', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))

    # --- Figure 6: Segmentation Output ---
    img_path_6 = os.path.join(IMG_DIR, 'result_02.png')
    if os.path.exists(img_path_6):
        img6_h = col_w / _img_ratio(img_path_6)
        story.append(Image(img_path_6, width=col_w, height=img6_h))
        story.append(Paragraph('Fig 6: Glacier Outlines & Segmentation Confidence Map', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))

    # --- Figure 7: Detailed Land Cover Map ---
    img_path_7 = os.path.join(IMG_DIR, '__results___30_0.png')
    if os.path.exists(img_path_7):
        img7_h = col_w / _img_ratio(img_path_7)
        story.append(Image(img_path_7, width=col_w, height=img7_h))
        story.append(Paragraph('Fig 7: Detailed Hyperspectral Land Cover Map (Tyrol)', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))
    
    # --- Figure 8: Class Distribution Analysis ---
    img_path_8 = os.path.join(IMG_DIR, '__results___21_0.png')
    if os.path.exists(img_path_8):
        img8_h = col_w / _img_ratio(img_path_8)
        story.append(Image(img_path_8, width=col_w, height=img8_h))
        story.append(Paragraph('Fig 8: Statistical Class Distribution Analysis', S('cap',fontSize=5.8,alignment=TA_CENTER)))
        story.append(SP(4))

    doc.build(story)
    print(f'PDF written -> {out}')
    
    # 2. Word Generation
    build_docx(out_docx)
    print(f'Word written -> {out_docx}')

if __name__ == '__main__':
    build()